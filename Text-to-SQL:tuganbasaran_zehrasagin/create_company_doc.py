import os
from docx import Document

def create_company_info_docx():
    """Creates a Word document with company information from the text file"""
    
    # Read the text file
    with open("data/NovaCart_RAG_Company_Info.txt", "r") as file:
        content = file.read()
    
    # Remove the triple quotes if present
    content = content.strip('"""')
    
    # Create a new Word document
    doc = Document()
    
    # Split the content by lines
    lines = content.split('\n')
    
    # Process each line
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if it's a heading
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            # Handle bullet points
            p = doc.add_paragraph()
            p.add_run(line[2:])
            p.style = 'List Bullet'
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save the document
    doc.save("data/NovaCart_RAG_Company_Info.docx")
    print("Company information document created successfully!")

# Run the function if this file is executed directly
if __name__ == "__main__":
    create_company_info_docx()